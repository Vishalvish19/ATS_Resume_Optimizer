import streamlit as st
import docx2txt
import PyPDF2
import re
import io
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document

st.set_page_config(page_title="ATS Resume Optimizer", layout="centered")

@st.cache_resource
def load_paraphraser():
    return pipeline("text2text-generation", model="mrm8488/t5-base-finetuned-question-generation-ap", tokenizer="t5-base")

paraphraser = load_paraphraser()

def read_resume(file):
    if file.name.endswith(".docx"):
        return docx2txt.process(file)
    elif file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        return " ".join(page.extract_text() for page in reader.pages if page.extract_text())
    return ""

def clean_text(text):
    text = re.sub(r"[^a-zA-Z0-9\s]", "", text.lower())
    return text

def extract_keywords(jd_text, top_n=15):
    words = re.findall(r"\b[a-zA-Z]{3,}\b", jd_text.lower())
    vectorizer = CountVectorizer(stop_words='english').fit([" ".join(words)])
    freqs = zip(vectorizer.get_feature_names_out(), vectorizer.transform([" ".join(words)]).toarray()[0])
    sorted_words = sorted(freqs, key=lambda x: x[1], reverse=True)
    return [word for word, freq in sorted_words[:top_n]]

def generate_summary(keywords):
    return f"Data Analyst with expertise in {', '.join(keywords[:5])}. Proven ability to drive insights using tools like {', '.join(keywords[5:8])}. Strong communication and cross-functional collaboration."

def rewrite_experience(resume_text, keywords):
    lines = resume_text.splitlines()
    rewritten = []
    for line in lines:
        if line.strip().startswith("-") or line.strip().startswith("\u2022"):
            prompt = f"paraphrase: {line.strip()} Include: {' '.join(keywords[:3])}"
            try:
                out = paraphraser(prompt, max_length=60, num_return_sequences=1)
                rewritten.append("\u2022 " + out[0]['generated_text'])
            except:
                rewritten.append(line)
        else:
            rewritten.append(line)
    return "\n".join(rewritten)

def update_skills(existing_skills, keywords):
    existing = set(map(str.lower, re.findall(r"\b\w+\b", existing_skills)))
    added = [kw for kw in keywords if kw.lower() not in existing]
    return existing_skills + ", " + ", ".join(added) if added else existing_skills

def build_resume(summary, experience, skills, original_text):
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Professional Experience", level=1)
    for line in experience.splitlines():
        if line.strip().startswith("\u2022"):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(skills)
    doc.add_heading("Education", level=1)
    for line in original_text.splitlines():
        if "education" in line.lower():
            doc.add_paragraph(line)
    return doc

def score_keywords(resume_text, jd_keywords):
    resume_words = set(clean_text(resume_text).split())
    matches = [kw for kw in jd_keywords if kw.lower() in resume_words]
    return len(matches), matches


st.title("🎯 ATS Resume Optimizer")
st.markdown("Upload your resume and job description to generate a keyword-aligned, ATS-friendly version.")

resume_file = st.file_uploader("📄 Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
jd_text = st.text_area("✏️ Paste the Job Description", height=300)
submit = st.button("🚀 Optimize Resume")

if submit:
    if resume_file and jd_text.strip():
        resume_text = read_resume(resume_file)
        jd_keywords = extract_keywords(jd_text)

        # Sections
        new_summary = generate_summary(jd_keywords)
        improved_experience = rewrite_experience(resume_text, jd_keywords)
        new_skills = update_skills("SQL, Python, Power BI", jd_keywords)

        
        keyword_count, matched_keywords = score_keywords(resume_text, jd_keywords)
        keyword_count_after, matched_keywords_after = score_keywords(new_summary + improved_experience + new_skills, jd_keywords)

      
        new_doc = build_resume(new_summary, improved_experience, new_skills, resume_text)
        buffer = io.BytesIO()
        new_doc.save(buffer)
        buffer.seek(0)

       
        st.subheader("📊 ATS Match Simulation")
        st.metric("Keyword Match Before", f"{keyword_count} / {len(jd_keywords)}")
        st.metric("Keyword Match After", f"{keyword_count_after} / {len(jd_keywords)}")

        st.subheader("✅ Keywords Matched After Optimization")
        st.markdown(", ".join(matched_keywords_after))

        st.download_button("📥 Download Optimized Resume", buffer, file_name="ATS_Optimized_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Please upload a resume and paste a job description.")
